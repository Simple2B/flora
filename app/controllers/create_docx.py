import os.path
import datetime
import uuid

from app.models import Bid, WorkItemGroup, LinkWorkItem
from app.controllers import calculate_alternate_total
from app.logger import log

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
PATH_TO_IMG = os.path.join(BASE_DIR, 'app/static/images/docx')

management_head_list = [
    ('DDB Contracting, LLC', 'management_style_submit'),
    ('Premier Project Management', 'management_style_approve')
]
management_info_list = [f'By{" "*10}', f'Name{" "*4}', f'Date{" "*6}']


def create_docx(bid_id):
    bid = Bid.query.get(bid_id)

    PATH_TO_SAVE_DOCX = f'/tmp/docx_{uuid.uuid4()}.docx'

    db_subtotal_data = {
        'Subtotal:': (bid.subtotal, 1),
        'Permit/Filing Free:': (bid.permit_filling_fee, bid.percent_permit_fee),
        'General Conditions:': (bid.general_conditions, bid.percent_general_condition),
        'Overhead:': (bid.insurance_tax, bid.percent_overhead),
        'Insurance/Tax:': (bid.overhead, bid.percent_insurance_tax),
        'Profit:': (bid.profit, bid.percent_profit),
        'Bond:': (bid.bond, bid.percent_bond),
        'Grand Total:': (bid.grand_subtotal, 1)
    }

    bid_global_work_items = LinkWorkItem.query.filter(
        LinkWorkItem.bid_id == bid_id).filter(LinkWorkItem.work_item_group == None).all()  # noqa 711
    groups = WorkItemGroup.query.filter(WorkItemGroup.bid_id == bid_id).all()

    def check_tbd(arg):
        if db_subtotal_data.get(arg, 'T.B.D')[0] == 0:
            return 'T.B.D'
        else:
            return db_subtotal_data.get(arg, 'T.B.D')[0]

    def set_row_height(row, arg, pt=True):
        row.height_rule = WD_ROW_HEIGHT.AT_LEAST
        if pt:
            row.height = Pt(arg)
        else:
            row.height = Cm(arg)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

    def write_to_docx(insert=False, cell_paragraph=None, edit_first_paragraph=False, content='', font_name='Arial',
                      font_size=12, font_bold=False, font_italic=False, font_underline=False, color=RGBColor(0, 0, 0),
                      font_highlight_color=None, before_spacing=5, after_spacing=5, line_spacing=1.34,
                      keep_together=True, keep_with_next=False, page_break_before=False, widow_control=False,
                      left_indent=None, right_indent=None, align='left', style=''):
        if insert:  # insert object into paragraph
            if cell_paragraph:
                font = cell_paragraph.add_run(content).font
                paragraph = cell_paragraph
                paragraph_format = paragraph.paragraph_format
            else:
                pass  # put here log_info
        else:
            if cell_paragraph:
                if edit_first_paragraph:
                    paragraph = cell_paragraph.paragraphs[0]
                    paragraph.style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
                    paragraph.text = str(content)
                else:
                    paragraph = cell_paragraph.add_paragraph(str(content))
            else:
                paragraph = document.add_paragraph(str(content))
                paragraph.style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)

            font = paragraph.style.font
            paragraph_format = paragraph.paragraph_format

        paragraph_format.left_indent = left_indent
        paragraph_format.right_indent = right_indent
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control

        if align.lower() == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align.lower() == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align.lower() == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align.lower() == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        font.name = font_name
        font.size = Pt(font_size)
        font.bold = font_bold
        font.italic = font_italic
        font.underline = font_underline
        font.highlight_color = font_highlight_color

        if not insert:
            font.color.rgb = color

        return paragraph

    # /// Document "margins to all document" block
    document = Document()
    sections = document.sections
    for section in sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1.34)
        section.bottom_margin = Cm(2.54)
        # section.header_distance = Mm(12.7)
        # section.footer_distance = Mm(12.7)

    # /// endblock

# ////// Begin to create document
# /// add bid information block
    document.add_picture(f'{PATH_TO_IMG}/logo_pdf.png', width=Cm(6.91), height=Cm(2.64))

    bid_table_info = document.add_table(rows=0, cols=4)
    bid_table_info.autofit = False
    bid_table_info.columns[0].width = Cm((9.5*0.2))
    bid_table_info.columns[1].width = Cm((9.5*0.8))
    bid_table_info.columns[2].width = Cm((9.5*0.35))
    bid_table_info.columns[3].width = Cm((9.5*0.65))    # entire width of 4 cols are 19,00 Cm

    row = bid_table_info.add_row()  # 1 row
    set_row_height(row, 15)
    cell_date = row.cells[3]
    write_to_docx(
        cell_paragraph=cell_date,
        edit_first_paragraph=True,
        content=f"{bid.due_date}",
        font_bold=True,
        font_size=10.5,
        align='right',
        style=f'bid_date_{"bid.date"}'
    )

    row = bid_table_info.add_row()  # 2 row with 'Client' and 'Project'
    set_row_height(row, 15)
    cell_client = row.cells[0]
    cell_client_name = row.cells[1]
    cell_project = row.cells[2]
    cell_project_name = row.cells[3]
    write_to_docx(
        cell_paragraph=cell_client,
        edit_first_paragraph=True,
        content="Client",
        font_bold=True,
        font_size=10.5,
        align='left',
        style='bid_client'
    )
    write_to_docx(
        cell_paragraph=cell_client_name,
        edit_first_paragraph=True,
        content=f"{bid.client}",
        font_bold=True,
        font_size=10.5,
        align='left',
        style=f'bid_client_{bid.client}'
    )
    write_to_docx(
        cell_paragraph=cell_project,
        edit_first_paragraph=True,
        content="Project",
        font_bold=True,
        font_size=10.5,
        align='left',
        left_indent=Cm(1),
        style='bid_project'
    )
    write_to_docx(
        cell_paragraph=cell_project_name,
        edit_first_paragraph=True,
        content=f"{bid.project_name}",
        font_bold=True,
        font_size=10.5,
        align='left',
        style='bid_client_Transaction Window & Sink'
    )

    row = bid_table_info.add_row()
    set_row_height(row, 15)
    cell_client_info = row.cells[1]
    cell_project_info = row.cells[3]
    write_to_docx(
        cell_paragraph=cell_client_info,
        edit_first_paragraph=True,
        content=f"{bid.vendor_address_street}",
        font_bold=False,
        align='left',
        style=f'bid_client_street_{bid.vendor_address_street}'
    )
    write_to_docx(
        cell_paragraph=cell_project_info,
        edit_first_paragraph=True,
        content=f"{bid.address_street}",
        font_bold=False,
        font_size=11.5,
        align='left',
        style=f'bid_project_street_{bid.address_street}'
    )

    # 4 row
    row = bid_table_info.add_row()
    set_row_height(row, 15)
    cell_client_info = row.cells[1]
    cell_project_info = row.cells[3]
    write_to_docx(
        cell_paragraph=cell_client_info,
        edit_first_paragraph=True,
        content=f"{bid.vendor_address_city}",
        font_bold=False,
        align='left',
        style=f'bid_client_{bid.vendor_address_city}'
    )
    write_to_docx(
        cell_paragraph=cell_project_info,
        edit_first_paragraph=True,
        content=f"{bid.address_city}",
        font_bold=False,
        font_size=10.5,
        align='left',
        style=f'bid_project_{bid.address_city}'
    )

    row = bid_table_info.add_row()
    set_row_height(row, 15)
    cell_client_info = row.cells[1]
    cell_project_info = row.cells[3]
    write_to_docx(
        cell_paragraph=cell_client_info,
        edit_first_paragraph=True,
        content=f"{bid.contact}",
        font_bold=True,
        font_size=10.5,
        align='left',
        style=f'bid_client_contact{bid.contact}'
    )
    write_to_docx(
        cell_paragraph=cell_project_info,
        edit_first_paragraph=True,
        content=f"{bid.project_type.value} # B-20-034 R{bid.revision}",
        font_bold=True,
        font_size=10.5,
        font_highlight_color=WD_COLOR_INDEX.YELLOW,
        align='left',
        style=f'bid_project_type_{bid.project_type.value} # B-20-034 R{bid.revision}'
    )

    # /// endblock
    # begin Section A (work_items) block
    document.add_paragraph()
    document.add_picture(f'{PATH_TO_IMG}/Section_A.png', width=Cm(18.99), height=Cm(0.65))
    write_to_docx(
        content='Please find our detailed Quote for the above referenced project as outlined below:',
        font_size=12.5,
        font_bold=True,
        after_spacing=10
    )

    #  / Render work_item's information
    work_item_table = document.add_table(rows=0, cols=3)
    work_item_table.columns[0].width = Cm(10.5)
    work_item_table.columns[1].width = Cm(5.695)
    work_item_table.columns[2].width = Cm(3.18)
    work_item_table.autofit = False

    if bid_global_work_items:
        for link_work_item in bid_global_work_items:
            row = work_item_table.add_row()
            row.height_rule = WD_ROW_HEIGHT.AT_LEAST
            row.height = Cm(0.55)
            paragraph_1 = row.cells[0].paragraphs[0]
            paragraph_2 = row.cells[2].paragraphs[0]

            write_to_docx(
                insert=True,
                cell_paragraph=paragraph_1,
                content=f'{link_work_item.work_item.code}',
                font_size=10.5,
                font_bold=True,
                style=f'bold_workitem_code_{link_work_item.work_item.code}'
            )
            paragraph_1.runs[0].add_tab()
            write_to_docx(
                insert=True,
                cell_paragraph=paragraph_1,
                content=f'{link_work_item.work_item.name}',
                font_size=10.5,
                font_bold=True,
                style=f'bold_workitem_name_{link_work_item.work_item.name}'
            )
            write_to_docx(
                insert=True,
                cell_paragraph=paragraph_2,
                content=(f'$ {link_work_item.link_subtotal}' if {link_work_item.link_subtotal} < 0.001 else 'TBD'),
                font_size=10.5,
                font_bold=True,
                align='right',
                style=f'bold_workitem_subtotal_{link_work_item.link_subtotal}'
            )
            for work_item_line in link_work_item.work_item_lines:
                work_item_line_table = document.add_table(rows=0, cols=3)
                work_item_line_table.columns[0].width = Cm(9.6875*0.25)
                work_item_line_table.columns[1].width = Cm(9.6875*0.5)
                work_item_line_table.columns[2].width = Cm(19.375/2 + 9.6875*0.25)
                work_item_line_table.autofit = False
                row = work_item_line_table.add_row()
                row.height_rule = WD_ROW_HEIGHT.AT_LEAST
                row.height = Cm(0.35)
                paragraph_1 = row.cells[1].paragraphs[0]
                paragraph_2 = row.cells[2].paragraphs[0]
                # second cell
                write_to_docx(
                    insert=True,
                    cell_paragraph=paragraph_1,
                    content=f'Note: {work_item_line.note}',
                    font_size=10.5,
                    style=f'work_item_line_note{work_item_line.id}'
                )
                write_to_docx(
                    cell_paragraph=row.cells[1],
                    content=f'{work_item_line.description}' + '(TBD)' if {work_item_line.tbd} else '',
                    font_size=9.5,
                    left_indent=Cm(0.5),
                    style=f'work_item_line_description{work_item_line.id}'
                )
                # third cell
                write_to_docx(
                    insert=True,
                    cell_paragraph=paragraph_2,
                    content=f'{work_item_line.quantity}     {work_item_line.unit}',
                    left_indent=Cm(2),
                    font_size=9.5,
                    style=f'work_item_line_q_u_{work_item_line.id}'
                )

    #  /  Render Work item group
    if groups:
        for group in groups:
            write_to_docx(
                content=f'{group.name}:',
                font_underline=True,
                font_bold=True,
                font_size=12,
                page_break_before=True,
                style=f'group_name_{group.name}'
            )
            document.add_picture(f'{PATH_TO_IMG}/group.png', width=Cm(19), height=Cm(0.65))

            group_work_item_table = document.add_table(rows=0, cols=3)
            group_work_item_table.columns[0].width = Cm(10.5)
            group_work_item_table.columns[1].width = Cm(5.695)
            group_work_item_table.columns[2].width = Cm(3.18)
            group_work_item_table.autofit = False

            for link_work_item in group.link_work_items:
                row = group_work_item_table.add_row()
                row.height_rule = WD_ROW_HEIGHT.AT_LEAST
                row.height = Cm(0.55)
                paragraph_1 = row.cells[0].paragraphs[0]
                paragraph_2 = row.cells[2].paragraphs[0]

                write_to_docx(
                    insert=True,
                    cell_paragraph=paragraph_1,
                    content=f'{link_work_item.work_item.code}',
                    font_size=10.5,
                    font_bold=True,
                    style='bold_workitem_id'
                )
                paragraph_1.runs[0].add_tab()
                write_to_docx(
                    insert=True,
                    cell_paragraph=paragraph_1,
                    content=f'{link_work_item.work_item.name}',
                    font_size=10.5,
                    font_bold=True,
                    style='bold_workitem_name'
                )
                write_to_docx(
                    insert=True,
                    cell_paragraph=paragraph_2,
                    content=(f'$ {link_work_item.link_subtotal}' if {link_work_item.link_subtotal} < 0.001 else 'TBD'),
                    font_size=10.5,
                    font_bold=True,
                    align='right',
                    style='bold_workitem_subtotal'
                )
                for work_item_line in link_work_item.work_item_lines:
                    group_work_item_line_table = document.add_table(rows=0, cols=3)
                    group_work_item_line_table.columns[0].width = Cm(9.6875*0.25)
                    group_work_item_line_table.columns[1].width = Cm(9.6875*0.5)
                    group_work_item_line_table.columns[2].width = Cm(19.375/2 + 9.6875*0.25)
                    group_work_item_line_table.autofit = False
                    row = group_work_item_line_table.add_row()
                    row.height_rule = WD_ROW_HEIGHT.AT_LEAST
                    row.height = Cm(0.35)
                    paragraph_1 = row.cells[1].paragraphs[0]
                    paragraph_2 = row.cells[2].paragraphs[0]
                    # second cell
                    write_to_docx(
                        insert=True,
                        cell_paragraph=paragraph_1,
                        content=f'Note: {work_item_line.note}',
                        font_size=10.5,
                        style=f'work_item_line_note_{work_item_line.id}'
                    )
                    write_to_docx(
                        cell_paragraph=row.cells[1],
                        content=f'{work_item_line.description}' + '(TBD)' if {work_item_line.tbd} else '',
                        font_size=9.5,
                        left_indent=Cm(0.5),
                        keep_with_next=True,
                        style=f'work_item_line_description_{work_item_line.id}'
                    )
                    # third cell
                    write_to_docx(
                        insert=True,
                        cell_paragraph=paragraph_2,
                        content=f'{work_item_line.quantity}',
                        font_size=9.5,
                        style=f'work_item_line_q_u_{work_item_line.id}'
                    )
                    write_to_docx(
                        insert=True,
                        cell_paragraph=paragraph_2,
                        content=f'   {work_item_line.unit}',
                        font_size=9.5,
                        left_indent=Cm(1.88),
                        style=f'work_item_line_q_u_{work_item_line.id}'
                    )

    #  / bid_subtotal on this section
    bid_subtotal_table = document.add_table(rows=0, cols=3)
    bid_subtotal_table.autofit = False
    bid_subtotal_table.columns[0].width = Cm(10.5)
    bid_subtotal_table.columns[1].width = Cm(5.695)
    bid_subtotal_table.columns[2].width = Cm(3.18)

    for i, j in enumerate(db_subtotal_data):
        if db_subtotal_data[j][1] < 0.001:
            log(log.DEBUG, "Percent of [%s]", str(db_subtotal_data[j]) + " " + str(db_subtotal_data[j][1]))
        else:
            row = bid_subtotal_table.add_row()
            set_row_height(row, 0.55, pt=False)
            cell_1 = row.cells[1]
            cell_1.width = Cm(5.695)  # 2050415
            cell_2 = row.cells[2]
            cell_2.width = Cm(3.18)
            write_to_docx(
                cell_paragraph=cell_1,
                edit_first_paragraph=True,
                content=j,
                font_bold=True,
                font_size=10.5,
                align='left',
                left_indent=Cm(1),
                style=f'bid_data_{j.lower()}'
            )
            write_to_docx(
                cell_paragraph=cell_2,
                edit_first_paragraph=True,
                content=(f'$ {check_tbd(j)}' if check_tbd(j) != 'T.B.D' else 'T.B.D'),
                font_bold=True,
                align='right',
                font_size=10.5,
                style=f'bid_data_{db_subtotal_data[j]}_{i}'
            )

# /// exclusions, clarifications, alternates blocks
    # begin Section B block
    document.add_picture(f'{PATH_TO_IMG}/Section_B.png', width=Cm(18.99), height=Cm(0.65))
    write_to_docx(
        content='Unless expressly stated, the following exclusions apply:',
        font_size=12.5,
        font_bold=True,
        after_spacing=10,
        style='exclusion_style_heading'
    )

    exclusion_paragraph = write_to_docx(
        content=' ',
        font_size=9.5,
        after_spacing=10,
        style='exclusion_style_first_paraprgaph'
    )
    if bid.exclusion_links:
        for exclusion_link in bid.exclusion_links:
            write_to_docx(
                insert=True,
                cell_paragraph=exclusion_paragraph,
                content=(f'{exclusion_link.exclusion.title}.' if exclusion_link == bid.exclusion_links[-1] else f'{exclusion_link.exclusion.title}, '),  # noqa 501
                font_size=9.5,
                after_spacing=10,
                style=f'exclusion_style_title_{i}'
            )
    # endblock

    # begin Section C block
    document.add_picture(f'{PATH_TO_IMG}/Section_C.png', width=Cm(18.99), height=Cm(0.65))
    write_to_docx(
        content='Please note the following clarifications:',
        font_size=12.5,
        font_bold=True,
        after_spacing=10,
        style='clarification_style_heading'
    )

    clarification_paragraph = write_to_docx(
        content=' ',
        font_size=9.5,
        after_spacing=10,
        style='clarification_style_first_paraprgaph'
    )
    if bid.clarification_links:
        for clarification_link in bid.clarification_links:
            write_to_docx(
                insert=True,
                cell_paragraph=clarification_paragraph,
                content=(f'{clarification_link.clarification.description}.' if clarification_link == bid.clarification_links[-1] else f'{clarification_link.clarification.description}, '),  # noqa 501
                font_size=9.5,
                style=f'clarification_style_title_{i}'
            )
    # endblock

    # document.add_page_break()

    # begin Section D block
    document.add_picture(f'{PATH_TO_IMG}/Section_D.png', width=Cm(18.99), height=Cm(0.65))

    alternate_table = document.add_table(rows=0, cols=3)
    alternate_table.autofit = False
    alternate_table.columns[0].width = Cm(16.195)
    alternate_table.columns[1].width = Cm(3.18)
    row = alternate_table.add_row()
    set_row_height(row, 15)
    alternate_paragraph = row.cells[0].paragraphs[0]
    paragraph_alternates_name = row.cells[0].paragraphs[0]
    paragraph_total_price = row.cells[1].paragraphs[0]
    if bid.alternates:
        for alternate in bid.alternates:
            write_to_docx(
                insert=True,
                cell_paragraph=paragraph_alternates_name,
                content=(f'{alternate.name}.' if alternate == bid.alternates[-1] else f'{alternate.name}, '),  # noqa 501
                font_size=9.5,
                style=f'alternate_style_title_{i}'
            )
        write_to_docx(
            insert=True,
            cell_paragraph=paragraph_total_price,
            content=f'$ {calculate_alternate_total(bid_id)}',
            font_bold=True,
            font_size=9.5,
            align='right',
            style=f'alternate_style_default_title_{i}'
        )
    else:
        write_to_docx(
                insert=True,
                cell_paragraph=alternate_paragraph,
                content='No alternates speciefied.',
                font_size=9.5,
                style=f'alternate_style_default_title_{i}'
            )
    # endblock

    document.add_page_break()
    # begin Section E_F block
    document.add_picture(f'{PATH_TO_IMG}/Section_E_F.png', width=Cm(18.99), height=Cm(0.79))

    management_table = document.add_table(rows=0, cols=2)
    management_table_row_head = management_table.add_row()
    for i, cell in enumerate(management_table_row_head.cells):
        write_to_docx(
            insert=True,
            content=management_head_list[i][0],
            cell_paragraph=cell.paragraphs[0],
            font_size=10.5,
            font_bold=True,
            align='center',
            after_spacing=20,
            style=management_head_list[i][1]
        )

    for i in range(3):
        management_table_row_info = management_table.add_row()
        for j, cell in enumerate(management_table_row_info.cells):
            paragraph = cell.paragraphs[0]
            write_to_docx(
                insert=True,
                content=management_info_list[i],
                cell_paragraph=paragraph,
                font_size=10,
                before_spacing=5,
                left_indent=Cm(1),
                # style='management'
            )
            paragraph.add_run()
            if i == 1 and j == 0:
                write_to_docx(
                    insert=True,
                    content=f'{bid.contact}',
                    cell_paragraph=paragraph,
                    font_size=10,
                    before_spacing=5,
                    left_indent=Cm(1)
                    # style='management'
                )
            elif i == 2 and j == 0:
                write_to_docx(
                    insert=True,
                    content=f'{datetime.datetime.today().strftime("%Y-%m-%d")}',
                    cell_paragraph=paragraph,
                    font_size=10,
                    before_spacing=5,
                    left_indent=Cm(1)
                    # style='management'
                )
            else:
                paragraph.runs[1].add_picture(f'{PATH_TO_IMG}/underline.png', width=Cm(3.75), height=Cm(0.05))
    # endblock
    document.add_paragraph()
    write_to_docx(
        content=f'{bid.project_type.value} # B_20_034 REVISED-{bid.revision}',
        font_size=10,
        after_spacing=10,
        style=f'quote {bid.project_type.value}'
    )

    document.save(PATH_TO_SAVE_DOCX)

    return PATH_TO_SAVE_DOCX
